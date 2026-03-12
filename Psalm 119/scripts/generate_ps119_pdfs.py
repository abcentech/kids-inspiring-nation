from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    import win32com.client  # type: ignore
except ImportError as exc:  # pragma: no cover
    raise SystemExit(
        "pywin32 is required to export DOCX files to PDF. Install it with 'python -m pip install pywin32'."
    ) from exc


ROOT = Path(r"C:\Users\ADMIN\Documents\KIN\Psalm 119")
SOURCE_DIR = ROOT / "Ps119 WorkBook"
DOCX_DIR = ROOT / "tmp" / "docs" / "docx"
PDF_DIR = ROOT / "downloads"

HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
NUMBERED_RE = re.compile(r"^\d+\.\s+")
INLINE_RE = re.compile(r"(\*\*.*?\*\*|\*.*?\*)")


def set_page_border(section) -> None:
    sect_pr = section._sectPr
    borders = sect_pr.first_child_found_in("w:pgBorders")
    if borders is None:
        borders = OxmlElement("w:pgBorders")
        borders.set(qn("w:offsetFrom"), "page")
        sect_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "18")
        element.set(qn("w:color"), "D9B44A")


def style_document(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        set_page_border(section)

    palette = (
        ("Title", 22, RGBColor(0x0D, 0x3B, 0x66)),
        ("Heading 1", 16, RGBColor(0x0D, 0x3B, 0x66)),
        ("Heading 2", 13, RGBColor(0x8C, 0x60, 0x18)),
        ("Heading 3", 11, RGBColor(0x0D, 0x3B, 0x66)),
    )
    for style_name, size, color in palette:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_runs(paragraph, text: str) -> None:
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_separator(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9B44A")
    border.append(bottom)
    p_pr.append(border)


def add_heading(doc: Document, level: int, text: str, first_block: bool, page_break_before: bool) -> None:
    if page_break_before:
        doc.add_page_break()

    if level == 1:
        paragraph = doc.add_paragraph(style="Title")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        paragraph = doc.add_paragraph(style="Heading 1")
    elif level == 3:
        paragraph = doc.add_paragraph(style="Heading 2")
    else:
        paragraph = doc.add_paragraph(style="Heading 3")

    add_runs(paragraph, text)
    if level == 1:
        paragraph.paragraph_format.space_after = Pt(6)
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(12 if not first_block else 0)


def build_docx(source_path: Path, output_path: Path) -> None:
    lines = source_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    style_document(doc)

    first_block = True
    current_section = None

    for raw_line in lines:
        line = raw_line.rstrip()
        if not line.strip():
            continue

        if line.strip() == "---":
            add_separator(doc)
            continue

        heading_match = HEADING_RE.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            page_break_before = level == 2 and current_section is not None and heading_text != "TABLE OF CONTENTS"
            add_heading(doc, level, heading_text, first_block, page_break_before)
            first_block = False
            if level == 2:
                current_section = heading_text
            continue

        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_runs(paragraph, line[2:].strip())
            continue

        if NUMBERED_RE.match(line):
            paragraph = doc.add_paragraph(style="Normal")
            add_runs(paragraph, line.strip())
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.first_line_indent = Inches(-0.15)
            paragraph.paragraph_format.space_after = Pt(3)
            continue

        paragraph = doc.add_paragraph(style="Normal")
        add_runs(paragraph, line)
        paragraph.paragraph_format.space_after = Pt(6)

        if line.startswith('"') and line.endswith('"'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def export_docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    word = win32com.client.DispatchEx("Word.Application")
    try:
        word.Visible = False
        document = word.Documents.Open(str(docx_path.resolve()))
        try:
            document.ExportAsFixedFormat(str(pdf_path.resolve()), 17)
        finally:
            document.Close(False)
    finally:
        word.Quit()


def main() -> None:
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_DIR.mkdir(parents=True, exist_ok=True)

    for source_path in sorted(SOURCE_DIR.glob("*.txt")):
        docx_path = DOCX_DIR / f"{source_path.stem}.docx"
        pdf_path = PDF_DIR / f"{source_path.stem}.pdf"
        build_docx(source_path, docx_path)
        export_docx_to_pdf(docx_path, pdf_path)
        print(f"Generated {pdf_path}")


if __name__ == "__main__":
    main()
